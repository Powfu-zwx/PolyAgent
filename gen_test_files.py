"""Generate test files for PolyAgent M5 manual integration testing."""

import os
from pathlib import Path


def main() -> None:
    output_dir = Path("test_uploads")
    output_dir.mkdir(exist_ok=True)

    # 1. test_short.txt — 200-500 Chinese characters
    txt_content = (
        "南京工业大学是一所以工为主、多学科协调发展的省属重点建设大学。"
        "学校坐落于历史文化名城南京，拥有丁家桥校区、虹桥校区和江浦校区。"
        "学校设有多个学院，涵盖工学、理学、管理学、经济学、文学、法学、艺术学等学科门类。"
        "近年来，学校在化学工程与技术、材料科学与工程等领域取得了显著的科研成果。"
        "学校积极推进国际交流与合作，与多所海外知名高校建立了合作关系。"
        "学校注重学生综合素质的培养，鼓励学生参加各类学术竞赛和社会实践活动。"
        "学校拥有完善的奖学金体系，包括国家奖学金、校级奖学金和企业奖学金等多种类型。"
        "学生可通过教务系统进行选课、查询成绩、办理各类手续等操作。"
        "学校图书馆藏书丰富，设有多个阅览室和自习室，为学生提供良好的学习环境。"
        "学校食堂提供多样化的餐饮服务，满足不同地区学生的饮食需求。"
    )
    (output_dir / "test_short.txt").write_text(txt_content, encoding="utf-8")
    print(f"[OK] test_short.txt ({len(txt_content)} chars)")

    # 2. test_empty.txt — 0 bytes
    (output_dir / "test_empty.txt").write_text("", encoding="utf-8")
    print("[OK] test_empty.txt (0 bytes)")

    # 3. test_note.md — markdown content
    md_content = """# PolyAgent 测试文档

## 项目简介

PolyAgent 是基于多 Agent 架构的政务与校园智能对话系统。

## 核心功能

- **知识问答**：回答政策法规、校园制度等事实性问题
- **摘要生成**：对长篇政策文件生成结构化摘要
- **公文写作**：草拟通知、申请书、公文等
- **办事引导**：以多轮对话方式引导用户完成事务流程

## 技术栈

系统采用 LangGraph 进行多 Agent 编排，使用 DeepSeek 作为主力大语言模型，
Chroma 作为向量数据库存储知识文档。

## 使用说明

用户通过自然语言输入需求，系统自动识别意图并路由至对应 Agent 完成任务。
支持单一意图和复合意图处理，例如可以同时完成信息查询和公文起草。
"""
    (output_dir / "test_note.md").write_text(md_content, encoding="utf-8")
    print(f"[OK] test_note.md ({len(md_content)} chars)")

    # 4. test_other.csv — for unsupported format test
    csv_content = """name,age,department
张三,22,计算机学院
李四,23,化工学院
王五,21,材料学院
"""
    (output_dir / "test_other.csv").write_text(csv_content, encoding="utf-8")
    print(f"[OK] test_other.csv ({len(csv_content)} chars)")

    # 5. test_doc.pdf — simple PDF with text
    pdf_text = (
        "This is a test PDF document for PolyAgent integration testing. "
        "The system should be able to extract this text using pdfplumber. "
        "PolyAgent supports knowledge QA, summary generation, document writing, "
        "and step-by-step guidance for administrative procedures."
    )
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 10, pdf_text)
        pdf.output(str(output_dir / "test_doc.pdf"))
        print("[OK] test_doc.pdf (via fpdf2)")
    except ImportError:
        # Fallback: create minimal valid PDF manually (ASCII only)
        _create_minimal_pdf(output_dir / "test_doc.pdf", pdf_text)
        print("[OK] test_doc.pdf (via minimal PDF fallback)")

    # 6. test_doc.docx — simple Word document
    try:
        from docx import Document

        doc = Document()
        doc.add_heading("PolyAgent 测试文档", level=1)
        doc.add_paragraph(
            "本文档用于测试 PolyAgent 系统的文件上传与摘要生成功能。"
        )
        doc.add_paragraph(
            "南京工业大学拥有完善的教学体系和科研平台，"
            "为学生提供了良好的学习和发展环境。"
            "学校积极推进产学研合作，在化工、材料等领域拥有多项国家级科研项目。"
        )
        doc.add_paragraph(
            "学生在校期间可以申请国家奖学金、校级奖学金等多种资助，"
            "同时学校提供勤工助学岗位，帮助家庭经济困难学生顺利完成学业。"
        )
        doc.save(str(output_dir / "test_doc.docx"))
        print("[OK] test_doc.docx (via python-docx)")
    except ImportError:
        print("[SKIP] test_doc.docx — python-docx not installed")

    print(f"\nAll files generated in: {output_dir.resolve()}")
    print("Please use these files for manual integration testing (D1-D6).")


def _create_minimal_pdf(path: Path, text: str) -> None:
    """Create a minimal valid PDF with ASCII text (fallback if fpdf2 unavailable)."""
    # Truncate to safe ASCII
    safe_text = text[:200]
    stream_content = f"BT /F1 12 Tf 50 700 Td ({safe_text}) Tj ET"
    stream_length = len(stream_content)

    pdf_bytes = f"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj

2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj

3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]
   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj

4 0 obj
<< /Length {stream_length} >>
stream
{stream_content}
endstream
endobj

5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj

xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000{350 + stream_length:03d} 00000 n 

trailer
<< /Size 6 /Root 1 0 R >>
startxref
0
%%EOF
"""
    path.write_text(pdf_bytes, encoding="ascii")


if __name__ == "__main__":
    main()
